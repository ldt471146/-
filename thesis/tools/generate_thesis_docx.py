from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT / "tmp" / "thesis" / "template_editable.docx"
MARKDOWN_PATH = ROOT / "thesis" / "draft" / "thesis_draft.md"
OUTPUT_DIR = ROOT / "outputs" / "thesis"

BODY_START_PARA_INDEX = 84
TITLE_PARA_INDEX = 69
AUTHOR_PARA_INDEX = 70
TUTOR_PARA_INDEX = 71

ABSTRACT_FIRST_SAMPLE = 72
ABSTRACT_BODY_SAMPLE = 73
ABSTRACT_LAST_SAMPLE = 74
ABSTRACT_KEYWORD_SAMPLE = 75
EN_TITLE_SAMPLE = 76
EN_AUTHOR_SAMPLE = 77
EN_TUTOR_SAMPLE = 78
EN_ABSTRACT_FIRST_SAMPLE = 79
EN_ABSTRACT_BODY_SAMPLE = 80
EN_KEYWORD_SAMPLE = 81
INTRO_SAMPLE = 84
BODY_SAMPLE = 89
REFERENCE_TITLE_SAMPLE = 310
REFERENCE_ITEM_SAMPLE = 311
ACK_TITLE_SAMPLE = 329
APPENDIX_TITLE_SAMPLE = 334
APPENDIX_ITEM_SAMPLE = 335
CONCLUSION_SAMPLE = 301

STYLE_ALIASES = {
    "标题 1": ["Heading 1", "Normal"],
    "标题 2": ["Heading 2", "Normal"],
    "标题 3": ["Heading 3", "Normal"],
    "题注": ["Caption", "Normal"],
}

HEADING_PREFIX_RE = re.compile(r"^\d+(?:\.\d+)*\s*")
LIST_PREFIX_RE = re.compile(r"^(?:\d+\.|-)\s*")
TABLE_TITLE_RE = re.compile(r"^表\d+(?:-\d+)?\s*")


@dataclass
class Block:
    kind: str
    value: str | list[list[str]]


def read_markdown() -> tuple[dict[str, str], list[Block]]:
    text = MARKDOWN_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    meta: dict[str, str] = {}
    start = 0
    if lines and lines[0].strip() == "---":
        end = 1
        while end < len(lines) and lines[end].strip() != "---":
            line = lines[end]
            if ":" in line:
                key, value = line.split(":", 1)
                meta[key.strip()] = value.strip()
            end += 1
        start = end + 1

    blocks: list[Block] = []
    index = start
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        if line.startswith("### "):
            blocks.append(Block("h3", line[4:].strip()))
            index += 1
            continue
        if line.startswith("## "):
            blocks.append(Block("h2", line[3:].strip()))
            index += 1
            continue
        if line.startswith("# "):
            blocks.append(Block("h1", line[2:].strip()))
            index += 1
            continue
        if line.startswith("![") and "](" in line and line.endswith(")"):
            caption = line[2 : line.index("](")]
            path = line[line.index("](") + 2 : -1]
            blocks.append(Block("image", f"{caption}|{path}"))
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            if len(rows) >= 2:
                rows.pop(1)
            blocks.append(Block("table", rows))
            continue

        if re.match(r"^\d+\.\s", line) or line.startswith("- "):
            blocks.append(Block("paragraph", line.strip()))
            index += 1
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if not next_line.strip() or next_line.startswith("#") or next_line.startswith("![") or next_line.startswith("|"):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        blocks.append(Block("paragraph", " ".join(paragraph_lines)))

    return meta, blocks


def paragraph_style(doc: Document, preferred: str, fallback: str = "Normal"):
    candidates = [preferred]
    candidates.extend(STYLE_ALIASES.get(preferred, []))
    if fallback not in candidates:
        candidates.append(fallback)
    for candidate in candidates:
        for style in doc.styles:
            if style.name == candidate:
                return style
    return doc.styles[0]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def replace_paragraph_text(paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            paragraph._p.remove(run._r)
        return
    paragraph.add_run(text)


def copy_paragraph_format(target, source) -> None:
    target.alignment = source.alignment
    target_format = target.paragraph_format
    source_format = source.paragraph_format
    target_format.left_indent = source_format.left_indent
    target_format.right_indent = source_format.right_indent
    target_format.first_line_indent = source_format.first_line_indent
    target_format.space_before = source_format.space_before
    target_format.space_after = source_format.space_after
    target_format.line_spacing = source_format.line_spacing
    target_format.line_spacing_rule = source_format.line_spacing_rule
    target_format.keep_together = source_format.keep_together
    target_format.keep_with_next = source_format.keep_with_next
    target_format.page_break_before = source_format.page_break_before
    target_format.widow_control = source_format.widow_control


def set_outline_level(paragraph, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), str(level))


def copy_numpr(paragraph, sample_paragraph) -> None:
    sample_properties = sample_paragraph._p.pPr
    if sample_properties is None or sample_properties.numPr is None:
        return
    paragraph_properties = paragraph._p.get_or_add_pPr()
    current_numpr = paragraph_properties.find(qn("w:numPr"))
    if current_numpr is not None:
        paragraph_properties.remove(current_numpr)
    paragraph_properties.append(deepcopy(sample_properties.numPr))


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_centered_line(paragraph, text: str, size: int) -> None:
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(size)


def strip_heading_prefix(text: str) -> str:
    return HEADING_PREFIX_RE.sub("", text).strip()


def split_abstract(blocks: list[Block]) -> tuple[list[str], str, list[str], str, list[Block]]:
    abstract_paragraphs: list[str] = []
    abstract_keywords = ""
    english_abstract_paragraphs: list[str] = []
    english_keywords = ""
    body_start = 0
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.kind == "h1" and strip_heading_prefix(str(block.value)) == "摘要":
            index += 1
            while index < len(blocks):
                current = blocks[index]
                if current.kind == "h1":
                    break
                if current.kind == "paragraph":
                    text = str(current.value).strip()
                    if text.startswith("关键词："):
                        abstract_keywords = text.split("：", 1)[1].strip()
                    else:
                        abstract_paragraphs.append(text)
                index += 1
            continue
        if block.kind == "h1" and strip_heading_prefix(str(block.value)) == "Abstract":
            index += 1
            while index < len(blocks):
                current = blocks[index]
                if current.kind == "h1":
                    break
                if current.kind == "paragraph":
                    text = str(current.value).strip()
                    lower_text = text.lower()
                    if lower_text.startswith("key words:") or lower_text.startswith("keywords:"):
                        english_keywords = text.split(":", 1)[1].strip()
                    else:
                        english_abstract_paragraphs.append(text)
                index += 1
            continue
        body_start = index
        break
    return abstract_paragraphs, abstract_keywords, english_abstract_paragraphs, english_keywords, blocks[body_start:]


def merge_paragraphs(paragraphs: list[str], max_count: int) -> list[str]:
    cleaned = [paragraph.strip() for paragraph in paragraphs if paragraph.strip()]
    if not cleaned:
        return [""] * max_count
    if len(cleaned) <= max_count:
        return cleaned + [""] * (max_count - len(cleaned))
    merged = cleaned[: max_count - 1]
    merged.append(" ".join(cleaned[max_count - 1 :]))
    return merged


def set_paragraph_text(paragraph, text: str) -> None:
    replace_paragraph_text(paragraph, text)


def populate_front_matter(
    doc: Document,
    meta: dict[str, str],
    abstract_paragraphs: list[str],
    abstract_keywords: str,
    english_abstract_paragraphs: list[str],
    english_keywords: str,
) -> None:
    cn_abstract = merge_paragraphs(abstract_paragraphs, 3)
    en_abstract = merge_paragraphs(english_abstract_paragraphs, 2)

    set_paragraph_text(doc.paragraphs[ABSTRACT_FIRST_SAMPLE - 1], f"摘要：{cn_abstract[0]}")
    set_paragraph_text(doc.paragraphs[ABSTRACT_BODY_SAMPLE - 1], cn_abstract[1])
    set_paragraph_text(doc.paragraphs[ABSTRACT_LAST_SAMPLE - 1], cn_abstract[2])
    set_paragraph_text(doc.paragraphs[ABSTRACT_KEYWORD_SAMPLE - 1], f"关键词：{abstract_keywords}")

    set_paragraph_text(doc.paragraphs[EN_TITLE_SAMPLE - 1], meta.get("title_en", "Title"))
    set_paragraph_text(
        doc.paragraphs[EN_AUTHOR_SAMPLE - 1],
        f"Student majoring in {meta.get('major_en', meta.get('major', 'Major'))}    {meta.get('author_en', meta.get('author', 'Name'))}",
    )
    set_paragraph_text(
        doc.paragraphs[EN_TUTOR_SAMPLE - 1],
        f"Tutor    {meta.get('tutor_en', meta.get('tutor', 'Tutor'))}",
    )
    set_paragraph_text(doc.paragraphs[EN_ABSTRACT_FIRST_SAMPLE - 1], f"Abstract: {en_abstract[0]}")
    set_paragraph_text(doc.paragraphs[EN_ABSTRACT_BODY_SAMPLE - 1], en_abstract[1])
    set_paragraph_text(doc.paragraphs[EN_KEYWORD_SAMPLE - 1], f"Key words: {english_keywords}")


def add_body_paragraph(doc: Document, sample_doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = paragraph_style(doc, "Normal")
    copy_paragraph_format(paragraph, sample_doc.paragraphs[BODY_SAMPLE - 1])
    if TABLE_TITLE_RE.match(text.strip()):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
    paragraph.add_run(text)


def add_reference_paragraph(doc: Document, sample_doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = paragraph_style(doc, "Normal")
    copy_paragraph_format(paragraph, sample_doc.paragraphs[REFERENCE_ITEM_SAMPLE - 1])
    copy_numpr(paragraph, sample_doc.paragraphs[REFERENCE_ITEM_SAMPLE - 1])
    paragraph.add_run(LIST_PREFIX_RE.sub("", text).strip())


def add_appendix_paragraph(doc: Document, sample_doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = paragraph_style(doc, "Normal")
    copy_paragraph_format(paragraph, sample_doc.paragraphs[APPENDIX_ITEM_SAMPLE - 1])
    paragraph.add_run(text)


def add_special_title(doc: Document, sample_doc: Document, title: str, sample_index: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = paragraph_style(doc, "Normal")
    copy_paragraph_format(paragraph, sample_doc.paragraphs[sample_index - 1])
    set_outline_level(paragraph, 0)
    run = paragraph.add_run(title)
    run.font.size = Pt(12)


def add_heading(doc: Document, sample_doc: Document, level: int, text: str) -> None:
    style_name = f"标题 {level}"
    paragraph = doc.add_paragraph()
    paragraph.style = paragraph_style(doc, style_name)
    clean_text = strip_heading_prefix(text.strip())
    if clean_text in {"引言", "结论与展望"}:
        sample_index = INTRO_SAMPLE if clean_text == "引言" else CONCLUSION_SAMPLE
        copy_numpr(paragraph, sample_doc.paragraphs[sample_index - 1])
    paragraph.add_run(clean_text)


def add_abstract(doc: Document, sample_doc: Document, paragraphs: list[str], keywords: str) -> None:
    if not paragraphs:
        return

    first = doc.add_paragraph()
    first.style = paragraph_style(doc, "Normal")
    copy_paragraph_format(first, sample_doc.paragraphs[ABSTRACT_FIRST_SAMPLE - 1])
    first.add_run("摘要：").bold = True
    first.add_run(paragraphs[0])

    for text in paragraphs[1:]:
        paragraph = doc.add_paragraph()
        paragraph.style = paragraph_style(doc, "Normal")
        copy_paragraph_format(paragraph, sample_doc.paragraphs[ABSTRACT_BODY_SAMPLE - 1])
        paragraph.add_run(text)

    if keywords:
        keyword_paragraph = doc.add_paragraph()
        keyword_paragraph.style = paragraph_style(doc, "Normal")
        copy_paragraph_format(keyword_paragraph, sample_doc.paragraphs[ABSTRACT_BODY_SAMPLE - 1])
        keyword_paragraph.add_run("关键词：").bold = True
        keyword_paragraph.add_run(keywords)


def render_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.add_paragraph()


def render_image(doc: Document, payload: str) -> None:
    caption, raw_path = payload.split("|", 1)
    image_path = ROOT / raw_path
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Inches(6.3))
    else:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"【图片缺失】{raw_path}")

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.style = paragraph_style(doc, "题注")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption)


def render_blocks(doc: Document, sample_doc: Document, blocks: Iterable[Block]) -> None:
    current_h1 = ""
    for block in blocks:
        if block.kind == "h1":
            current_h1 = strip_heading_prefix(str(block.value))
            if current_h1 == "绪论":
                continue
            if current_h1 == "参考文献":
                add_special_title(doc, sample_doc, "参考文献", REFERENCE_TITLE_SAMPLE)
            elif current_h1 == "致谢":
                add_special_title(doc, sample_doc, "致谢", ACK_TITLE_SAMPLE)
            elif current_h1 == "附录":
                add_special_title(doc, sample_doc, "附录", APPENDIX_TITLE_SAMPLE)
            else:
                add_heading(doc, sample_doc, 1, str(block.value))
            continue
        if block.kind == "h2":
            add_heading(doc, sample_doc, 2, str(block.value))
            continue
        if block.kind == "h3":
            add_heading(doc, sample_doc, 3, str(block.value))
            continue
        if block.kind == "table":
            render_table(doc, block.value)  # type: ignore[arg-type]
            continue
        if block.kind == "image":
            render_image(doc, str(block.value))
            continue

        text = str(block.value)
        if current_h1 == "参考文献":
            add_reference_paragraph(doc, sample_doc, text)
        elif current_h1 == "附录":
            add_appendix_paragraph(doc, sample_doc, text)
        else:
            add_body_paragraph(doc, sample_doc, text)


def trim_template_body(doc: Document) -> None:
    for paragraph in list(doc.paragraphs)[BODY_START_PARA_INDEX - 1 :]:
        delete_paragraph(paragraph)


def restore_body_section(doc: Document, sample_doc: Document) -> None:
    body = doc._element.body
    sample_body = sample_doc._element.body
    current_sectpr = body.sectPr
    if current_sectpr is not None:
        body.remove(current_sectpr)
    sample_sectpr = sample_body.sectPr
    if sample_sectpr is not None:
        body.append(deepcopy(sample_sectpr))


def update_headers(doc: Document, title: str) -> None:
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip() == "论文题目":
                replace_paragraph_text(paragraph, title)


def build_output_path(title: str) -> Path:
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title).strip() or "论文题目"
    return OUTPUT_DIR / f"{safe_title}_修订稿.docx"


def ensure_writable_output_path(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with open(path, "ab"):
            return path
    except PermissionError:
        stem = path.stem
        suffix = path.suffix
        for index in range(2, 100):
            candidate = path.with_name(f"{stem}_{index}{suffix}")
            if not candidate.exists():
                return candidate
            try:
                with open(candidate, "ab"):
                    return candidate
            except PermissionError:
                continue
    return path


def main() -> None:
    meta, blocks = read_markdown()
    template = Document(str(TEMPLATE_PATH))
    sample_doc = Document(str(TEMPLATE_PATH))
    title_cn = meta.get("title_cn", "论文题目")

    enable_update_fields_on_open(template)
    set_centered_line(template.paragraphs[TITLE_PARA_INDEX - 1], title_cn, 16)
    set_centered_line(
        template.paragraphs[AUTHOR_PARA_INDEX - 1],
        f"{meta.get('major', '【专业待补】')}    {meta.get('author', '【姓名待补】')}",
        12,
    )
    set_centered_line(
        template.paragraphs[TUTOR_PARA_INDEX - 1],
        f"指导教师    {meta.get('tutor', '【指导教师待补】')}",
        12,
    )
    update_headers(template, title_cn)
    abstract_paragraphs, abstract_keywords, english_abstract_paragraphs, english_keywords, body_blocks = split_abstract(blocks)
    populate_front_matter(
        template,
        meta,
        abstract_paragraphs,
        abstract_keywords,
        english_abstract_paragraphs,
        english_keywords,
    )

    trim_template_body(template)
    render_blocks(template, sample_doc, body_blocks)
    restore_body_section(template, sample_doc)
    update_headers(template, title_cn)

    output_path = ensure_writable_output_path(build_output_path(title_cn))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    template.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
