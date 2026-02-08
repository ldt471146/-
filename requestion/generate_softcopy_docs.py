import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
REQ_DIR = ROOT / "requestion"
BACK_DIR = ROOT / "back" / "src"
FRONT_DIR = ROOT / "front" / "src"

SOFTWARE_NAME = "青少年编程平台"
VERSION = "1.0最终"


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
    style.font.size = Pt(10.5)
    fmt = style.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def set_source_doc_style(doc: Document) -> None:
    set_doc_style(doc)
    style = doc.styles["Normal"]
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(13)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_text(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def find_images() -> list[Path]:
    images = []
    for p in sorted(REQ_DIR.glob("*.png")):
        if p.name.lower().startswith("tmp"):
            continue
        images.append(p)
    for p in sorted(REQ_DIR.glob("*.jpg")):
        images.append(p)
    for p in sorted(REQ_DIR.glob("*.jpeg")):
        images.append(p)
    return images


def fit_image_size(image_path: Path, max_width_in: float, max_height_in: float) -> tuple[Inches, Inches]:
    with Image.open(image_path) as im:
        w, h = im.size
    if w <= 0 or h <= 0:
        return Inches(max_width_in), Inches(max_height_in)
    ratio = min(max_width_in / w, max_height_in / h)
    return Inches(w * ratio), Inches(h * ratio)


def build_operation_doc(images: list[Path]) -> Path:
    doc = Document()
    set_doc_style(doc)

    add_heading(doc, f"{SOFTWARE_NAME}V{VERSION}操作说明书")
    add_text(doc, "")
    add_text(doc, f"文档生成日期：{datetime.now().strftime('%Y-%m-%d')}")

    add_heading(doc, "一、引言")
    add_text(doc, "本文档用于说明本系统的核心操作方式、功能入口和常见异常处理方法，便于用户快速上手和日常使用。")

    add_heading(doc, "二、系统概述及用途")
    add_text(doc, "本系统面向青少年编程学习场景，提供课程学习、题库练习、成长报告、消息通知和教师审核等能力。")
    add_text(doc, "系统支持多角色协同，覆盖学生学习、教师内容发布与管理端审核流程。")

    add_heading(doc, "三、运行环境")
    add_text(doc, "1. 客户端：支持常见现代浏览器。")
    add_text(doc, "2. 服务端：支持常见服务器操作系统。")
    add_text(doc, "3. 数据层：关系型数据库、缓存服务、消息队列。")
    add_text(doc, "4. 网络要求：客户端可访问应用服务地址。")

    add_heading(doc, "四、功能模块说明")
    modules = [
        "1. 账号与认证模块：支持邮箱登录、注册、验证码校验、记住我。",
        "2. 学习中心模块：支持课程浏览、课程详情、课时进度、最近学习记录。",
        "3. 题库练习模块：支持按课程筛选、分页、错题本与收藏管理。",
        "4. 消息通知模块：支持分页查看、已读处理、删除消息。",
        "5. 教师与管理模块：支持教师申请、审核、状态回执。",
        "6. 智能助手模块：支持悬浮入口、会话问答。",
    ]
    for m in modules:
        add_text(doc, m)

    add_text(doc, "以下为界面截图（含图号占位）：")
    for idx, image in enumerate(images, start=1):
        cap = doc.add_paragraph(f"图{idx}  {image.stem}")
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try:
            # A4 可视区域内自适应，避免截图被裁切或越界
            pic_w, pic_h = fit_image_size(image, max_width_in=5.8, max_height_in=7.6)
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = None
            p.add_run().add_picture(str(image), width=pic_w, height=pic_h)
        except Exception:
            add_text(doc, f"[图{idx}截图占位：{image.name}]")
        add_text(doc, "")

    add_heading(doc, "五、主要操作流程")
    add_text(doc, "1. 用户注册：输入邮箱、用户名、密码、确认密码和验证码，提交完成注册。")
    add_text(doc, "2. 用户登录：输入邮箱和密码，可勾选记住我后登录。")
    add_text(doc, "3. 课程学习：进入课程列表，选择课程后进入详情页并继续学习课时。")
    add_text(doc, "4. 题目练习：进入题库，选择题目作答并查看解析，错题进入错题本。")
    add_text(doc, "5. 教师申请：用户提交教师申请后，等待管理端审核并接收结果通知。")

    add_heading(doc, "六、异常处理")
    add_text(doc, "1. 登录失败：检查输入内容格式与账号状态。")
    add_text(doc, "2. 验证码发送失败：检查邮箱地址有效性和发送频率限制。")
    add_text(doc, "3. 页面空白或数据缺失：检查服务端是否运行及网络连通性。")
    add_text(doc, "4. 权限不足：根据提示切换账号或联系管理员处理。")

    add_heading(doc, "七、附录")
    add_text(doc, "附录A：术语说明")
    add_text(doc, "1. 课时：课程中的最小学习单元。")
    add_text(doc, "2. 错题本：用户练习中答错题目的集中管理区。")
    add_text(doc, "3. 审核：管理端对教师申请状态的处理。")

    out = REQ_DIR / f"{SOFTWARE_NAME}V{VERSION}操作说明书.docx"
    doc.save(out)
    return out


SENSITIVE_PATTERNS = [
    (re.compile(r"(api[-_ ]?key\s*[:=]\s*)([^\s\"']+)", re.IGNORECASE), r"\1***"),
    (re.compile(r"(token\s*[:=]\s*)([^\s\"']+)", re.IGNORECASE), r"\1***"),
    (re.compile(r"(password\s*[:=]\s*)([^\s\"']+)", re.IGNORECASE), r"\1***"),
    (re.compile(r"(secret\s*[:=]\s*)([^\s\"']+)", re.IGNORECASE), r"\1***"),
    (re.compile(r"(username\s*[:=]\s*)([^\s\"']+)", re.IGNORECASE), r"\1***"),
    (re.compile(r"(ak|sk)-[A-Za-z0-9\-_]{10,}", re.IGNORECASE), "***"),
]


def sanitize(line: str) -> str:
    s = line.rstrip("\n")
    for pattern, repl in SENSITIVE_PATTERNS:
        s = pattern.sub(repl, s)
    return s


def format_code_line(line: str, max_len: int = 72) -> str:
    s = sanitize(line)
    # 防止 Word 自动换行导致页数失控，超长行做截断
    if len(s) > max_len:
        return s[: max_len - 3] + "..."
    return s


def collect_source_lines() -> list[str]:
    exts = {".java", ".xml", ".yml", ".yaml", ".js", ".ts", ".vue", ".css", ".scss", ".html"}
    roots = [BACK_DIR, FRONT_DIR]
    all_lines: list[str] = []
    for root in roots:
        if not root.exists():
            continue
        for file in sorted(root.rglob("*")):
            if not file.is_file():
                continue
            if file.suffix.lower() not in exts:
                continue
            rel = file.relative_to(ROOT).as_posix()
            all_lines.append(format_code_line(f"// ===== FILE: {rel} ====="))
            try:
                text = file.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                text = file.read_text(encoding="utf-8", errors="ignore")
            for idx, line in enumerate(text.splitlines(), start=1):
                all_lines.append(format_code_line(f"{idx:05d}: {line}"))
            all_lines.append("")
    return all_lines


def select_lines_for_doc(lines: list[str]) -> list[str]:
    if not lines:
        return ["// 暂无可导出的源代码"]
    if len(lines) > 3000:
        head = lines[:1500]
        tail = lines[-1500:]
        return head + ["", "// ===== 中间部分按要求省略 =====", ""] + tail
    return lines


def build_source_doc(lines: list[str]) -> Path:
    doc = Document()
    set_source_doc_style(doc)
    add_heading(doc, f"{SOFTWARE_NAME}V{VERSION}源代码（节选）")
    add_text(doc, "说明：本文件按软著文档要求输出，已对密钥、口令、令牌等敏感信息进行过滤处理。")
    add_text(doc, "")
    selected = select_lines_for_doc(lines)
    for line in selected:
        add_text(doc, line)
    out = REQ_DIR / f"{SOFTWARE_NAME}V{VERSION}源代码（节选）.docx"
    doc.save(out)
    return out


def build_source_front_doc(lines: list[str]) -> Path:
    doc = Document()
    set_source_doc_style(doc)
    # 严格按 30 页生成：每页 50 行，共 1500 行
    front_lines = lines[:1500] if len(lines) >= 1500 else lines
    lines_per_page = 50
    total_pages = 30
    needed = lines_per_page * total_pages
    if len(front_lines) < needed:
        front_lines = front_lines + [""] * (needed - len(front_lines))
    else:
        front_lines = front_lines[:needed]
    for page in range(total_pages):
        chunk = front_lines[page * lines_per_page : (page + 1) * lines_per_page]
        for line in chunk:
            add_text(doc, line)
        if page < total_pages - 1:
            doc.add_page_break()
    out = REQ_DIR / f"{SOFTWARE_NAME}V{VERSION}源代码（前30页）.docx"
    doc.save(out)
    return out


def build_source_tail_doc(lines: list[str]) -> Path:
    doc = Document()
    set_source_doc_style(doc)
    # 严格按 30 页生成：每页 50 行，共 1500 行
    tail_lines = lines[-1500:] if len(lines) >= 1500 else lines
    lines_per_page = 50
    total_pages = 30
    needed = lines_per_page * total_pages
    if len(tail_lines) < needed:
        tail_lines = tail_lines + [""] * (needed - len(tail_lines))
    else:
        tail_lines = tail_lines[:needed]
    for page in range(total_pages):
        chunk = tail_lines[page * lines_per_page : (page + 1) * lines_per_page]
        for line in chunk:
            add_text(doc, line)
        if page < total_pages - 1:
            doc.add_page_break()
    out = REQ_DIR / f"{SOFTWARE_NAME}V{VERSION}源代码（后30页）.docx"
    doc.save(out)
    return out


def main() -> None:
    images = find_images()
    op_file = build_operation_doc(images)
    source_lines = collect_source_lines()
    src_front_file = build_source_front_doc(source_lines)
    src_tail_file = build_source_tail_doc(source_lines)
    print(str(op_file))
    print(str(src_front_file))
    print(str(src_tail_file))
    print(f"IMAGES_USED={len(images)}")
    print(f"SOURCE_LINES={len(source_lines)}")


if __name__ == "__main__":
    main()
